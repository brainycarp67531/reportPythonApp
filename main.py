import tkinter as tk
from tkinter import messagebox
from docx import Document
import fitz  # PyMuPDF

def generate_pdf():
    material = material_entry.get()
    weight = weight_entry.get()

    if not material or not weight:
        messagebox.showerror("Input Error", "Please enter both Material and Weight.")
        return

    try:
        # Load the Word template
        doc = Document("template.docx")

        # Replace placeholders
        for para in doc.paragraphs:
            if "{{material}}" in para.text:
                para.text = para.text.replace("{{material}}", material)
            if "{{weight}}" in para.text:
                para.text = para.text.replace("{{weight}}", weight)

        # Save updated document
        doc.save("output.docx")

        # Convert to PDF using PyMuPDF
        pdf_doc = fitz.open()
        for para in doc.paragraphs:
            page = pdf_doc.new_page()
            page.insert_text((72, 72), para.text, fontsize=12)

        pdf_doc.save("output.pdf")
        pdf_doc.close()

        messagebox.showinfo("Success", "PDF generated successfully as 'output.pdf'.")

    except Exception as e:
        messagebox.showerror("Error", f"An error occurred: {e}")

# GUI setup
root = tk.Tk()
root.title("Material and Weight PDF Generator")

tk.Label(root, text="Material:").grid(row=0, column=0, padx=10, pady=5, sticky="e")
material_entry = tk.Entry(root, width=30)
material_entry.grid(row=0, column=1, padx=10, pady=5)

tk.Label(root, text="Weight:").grid(row=1, column=0, padx=10, pady=5, sticky="e")
weight_entry = tk.Entry(root, width=30)
weight_entry.grid(row=1, column=1, padx=10, pady=5)

generate_button = tk.Button(root, text="Generate PDF", command=generate_pdf)
generate_button.grid(row=2, column=0, columnspan=2, pady=10)

root.mainloop()
